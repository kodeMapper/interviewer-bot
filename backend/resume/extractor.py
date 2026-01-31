"""
Resume Text Extractor
Extracts raw text from resume files (PDF, DOCX, TXT)

Supports:
- PDF: Direct text extraction with PyMuPDF (fitz)
- DOCX: python-docx library
- TXT: Plain text reading
- OCR fallback for scanned PDFs (optional, requires Tesseract)
"""

import os
from typing import Tuple, Optional
from pathlib import Path


class ResumeExtractor:
    """
    Extracts text from resume files.
    Returns tuple: (extracted_text, extraction_method, confidence_score)
    """
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc', '.txt']
    
    def __init__(self):
        self.ocr_available = self._check_ocr_availability()
        self._fitz_available = self._check_fitz_availability()
        self._docx_available = self._check_docx_availability()
    
    def extract(self, file_path: str) -> Tuple[str, str, float]:
        """
        Main extraction method.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Tuple of (extracted_text, method_used, confidence_0_to_1)
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {ext}. Supported: {self.SUPPORTED_FORMATS}")
        
        if ext == '.pdf':
            return self._extract_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._extract_docx(file_path)
        elif ext == '.txt':
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported format: {ext}")
    
    def _extract_pdf(self, path: str) -> Tuple[str, str, float]:
        """
        PDF extraction with multi-layer fallback:
        1. Direct text extraction (fastest) using PyMuPDF
        2. OCR fallback (for scanned documents)
        """
        if not self._fitz_available:
            raise ImportError("PyMuPDF (fitz) not installed. Run: pip install PyMuPDF")
        
        import fitz  # PyMuPDF
        
        try:
            doc = fitz.open(path)
            text = ""
            
            for page in doc:
                text += page.get_text()
            
            doc.close()
            
            # Clean up text
            text = self._clean_text(text)
            
            # If text is too short, might be scanned - try OCR
            if len(text.strip()) < 100:
                if self.ocr_available:
                    print("   [Extractor] PDF has little text, trying OCR...")
                    return self._ocr_extract_pdf(path)
                else:
                    print("   [Extractor] Warning: PDF may be scanned. OCR not available.")
            
            # Calculate confidence based on text length and quality
            confidence = self._calculate_confidence(text)
            return text.strip(), "pdf_direct", confidence
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract PDF: {e}")
    
    def _extract_docx(self, path: str) -> Tuple[str, str, float]:
        """Extract from DOCX using python-docx."""
        if not self._docx_available:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        from docx import Document
        
        try:
            doc = Document(path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Also extract from tables (resumes often use tables for layout)
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))
            
            # Combine all text
            all_text = paragraphs + table_text
            text = "\n".join(all_text)
            text = self._clean_text(text)
            
            confidence = self._calculate_confidence(text)
            return text.strip(), "docx", confidence
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract DOCX: {e}")
    
    def _extract_txt(self, path: str) -> Tuple[str, str, float]:
        """Extract from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as f:
                        text = f.read()
                    text = self._clean_text(text)
                    confidence = self._calculate_confidence(text)
                    return text.strip(), "txt", confidence
                except UnicodeDecodeError:
                    continue
            
            raise RuntimeError("Could not decode text file with any known encoding")
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract TXT: {e}")
    
    def _ocr_extract_pdf(self, path: str) -> Tuple[str, str, float]:
        """OCR fallback for scanned PDFs using pdf2image + pytesseract."""
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            # Convert PDF pages to images
            images = convert_from_path(path, dpi=300)
            
            text_parts = []
            for i, image in enumerate(images):
                page_text = pytesseract.image_to_string(image)
                text_parts.append(page_text)
            
            text = "\n".join(text_parts)
            text = self._clean_text(text)
            
            # OCR confidence is lower
            confidence = min(0.7, self._calculate_confidence(text) * 0.8)
            return text.strip(), "pdf_ocr", confidence
            
        except Exception as e:
            raise RuntimeError(f"OCR extraction failed: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        import re
        
        # Replace multiple spaces with single space
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Replace multiple newlines with double newline
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text
    
    def _calculate_confidence(self, text: str) -> float:
        """
        Calculate extraction confidence based on text quality.
        
        Factors:
        - Text length (longer = more confident)
        - Presence of common resume keywords
        - Character quality (not too many special chars)
        """
        if not text:
            return 0.0
        
        # Base confidence on length
        length = len(text)
        if length < 100:
            length_score = 0.2
        elif length < 500:
            length_score = 0.5
        elif length < 1500:
            length_score = 0.8
        else:
            length_score = 1.0
        
        # Check for resume keywords
        resume_keywords = [
            'experience', 'education', 'skills', 'project', 'work',
            'internship', 'university', 'degree', 'certificate',
            'email', 'phone', 'address', 'linkedin', 'github'
        ]
        
        text_lower = text.lower()
        keyword_count = sum(1 for kw in resume_keywords if kw in text_lower)
        keyword_score = min(1.0, keyword_count / 5)  # Expect at least 5 keywords
        
        # Calculate final confidence
        confidence = (length_score * 0.6) + (keyword_score * 0.4)
        return round(confidence, 2)
    
    def _check_fitz_availability(self) -> bool:
        """Check if PyMuPDF is available."""
        try:
            import fitz
            return True
        except ImportError:
            return False
    
    def _check_docx_availability(self) -> bool:
        """Check if python-docx is available."""
        try:
            from docx import Document
            return True
        except ImportError:
            return False
    
    def _check_ocr_availability(self) -> bool:
        """Check if Tesseract OCR is installed."""
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            return True
        except:
            return False
    
    def get_supported_formats(self) -> list:
        """Return list of supported file formats."""
        return self.SUPPORTED_FORMATS.copy()
    
    def is_supported(self, file_path: str) -> bool:
        """Check if a file format is supported."""
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.SUPPORTED_FORMATS


# Standalone usage
if __name__ == "__main__":
    extractor = ResumeExtractor()
    
    # Test with a sample file
    import sys
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        text, method, confidence = extractor.extract(file_path)
        print(f"Method: {method}")
        print(f"Confidence: {confidence}")
        print(f"Text length: {len(text)}")
        print("\n--- Extracted Text ---")
        print(text[:1000] + "..." if len(text) > 1000 else text)
